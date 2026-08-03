from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Guia_completa_de_uso_Expediente_Integro.docx"
PDF = ROOT / "docs" / "Guia_completa_de_uso_Expediente_Integro.pdf"

document = Document(DOCX)
doc_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
doc_text += "\n" + "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)

reader = PdfReader(PDF)
pdf_text = "\n".join((page.extract_text() or "") for page in reader.pages)
headings = [
    paragraph.text
    for paragraph in document.paragraphs
    if paragraph.style.name in {"Heading 1", "Heading 2"}
]
with ZipFile(DOCX) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")

bad_tokens = ["�", "Lorem ipsum", "TODO", "replace_with", "{{", "}}"]
result = {
    "doc_paragraphs": len(document.paragraphs),
    "tables": len(document.tables),
    "headings": len(headings),
    "pdf_pages": len(reader.pages),
    "doc_chars": len(doc_text),
    "pdf_chars": len(pdf_text),
    "alt_text_entries": document_xml.count('descr="Figura'),
    "semantic_table_headers": document_xml.count("tblHeader"),
    "bad_tokens": [token for token in bad_tokens if token in doc_text or token in pdf_text],
    "has_title": "Expediente Íntegro" in doc_text,
    "has_dev_password": "2318" in doc_text,
    "has_role_table": (
        "Administración TI" in doc_text
        and "web.juez.20260731@example.test" in doc_text
    ),
    "has_final_marker": "Fin de la guía" in doc_text,
}

print(result)

assert result["pdf_pages"] == 25
assert result["alt_text_entries"] == 2
assert result["semantic_table_headers"] == 12
assert not result["bad_tokens"]
assert result["has_title"]
assert result["has_dev_password"]
assert result["has_role_table"]
assert result["has_final_marker"]
