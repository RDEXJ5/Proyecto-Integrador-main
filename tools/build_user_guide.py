from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
ASSET_DIR = OUTPUT_DIR / "_guide_assets"
OUTPUT_PATH = OUTPUT_DIR / "Guia_completa_de_uso_Expediente_Integro.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

# compact_reference_guide + named "Expediente Íntegro" brand override.
NAVY = "0D2740"
TEAL = "0B7A75"
DARK_TEAL = "075B58"
PALE_TEAL = "DDF3EF"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F3F6FA"
MID_GRAY = "64748B"
DARK = "162033"
WHITE = "FFFFFF"
GOLD = "B7791F"
RED = "A83232"
PALE_RED = "FDECEC"
PALE_GOLD = "FFF4D6"
GREEN = "247A48"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D8E0EA", size=6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)


def set_run_font(run, name="Calibri", size=11, color=DARK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MID_GRAY)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = p_pr.find(qn("w:keepLines"))
    if keep_lines is None:
        p_pr.append(OxmlElement("w:keepLines"))


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, TEAL, 14, 7),
        ("Heading 3", 12, DARK_TEAL, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    doc.styles["Heading 1"].paragraph_format.page_break_before = False

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ("List Bullet 2", "List Number 2"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.7)
        style.paragraph_format.first_line_indent = Inches(-0.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.2

    if "Guide Code" not in doc.styles:
        code_style = doc.styles.add_style("Guide Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor.from_string(DARK)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.18)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(8)
        code_style.paragraph_format.line_spacing = 1.1

    if "Guide Caption" not in doc.styles:
        caption_style = doc.styles.add_style("Guide Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Calibri"
        caption_style.font.size = Pt(9)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor.from_string(MID_GRAY)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_before = Pt(3)
        caption_style.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("EXPEDIENTE ÍNTEGRO  |  GUÍA DE USUARIO")
    set_run_font(run, size=8.5, color=TEAL, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Uso interno · Desarrollo y pruebas  |  Página ")
    set_run_font(run, size=9, color=MID_GRAY)
    add_page_field(fp)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GESTIÓN DOCUMENTAL JURÍDICA")
    set_run_font(run, size=10, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Expediente Íntegro")
    set_run_font(run, size=31, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Guía completa de uso, operación y pruebas")
    set_run_font(run, size=15, color=DARK_TEAL)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [CONTENT_WIDTH_DXA])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_border(cell, color="B6DDD8", size=7)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Manual para usuarios web, participantes móviles y personal técnico")
    set_run_font(r, size=11, color=NAVY, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(22)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2700, 6660])
    metadata = [
        ("Versión", "1.0"),
        ("Fecha", "1 de agosto de 2026"),
        ("Entorno", "Desarrollo y validación funcional"),
        ("Cobertura", "Web Flask, API Node.js, MySQL, MinIO y aplicación móvil"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_border(row.cells[0])
        set_cell_border(row.cells[1])
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, size=9.5, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de referencia para demostraciones, capacitación y pruebas de aceptación")
    set_run_font(r, size=9.5, color=MID_GRAY, italic=True)
    doc.add_page_break()


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    return p


def add_para(doc, text: str, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    set_keep_together(p)
    return p


def add_bullets(doc, items: list[str], level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for index, item in enumerate(items):
        p = doc.add_paragraph(style=style)
        p.add_run(item)
        set_keep_together(p)
        if index < len(items) - 1:
            p.paragraph_format.keep_with_next = True


def create_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_num_id = "0"
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        level = abstract_num.find(qn("w:lvl"))
        if level is None or level.get(qn("w:ilvl")) != "0":
            continue
        number_format = level.find(qn("w:numFmt"))
        if number_format is not None and number_format.get(qn("w:val")) == "decimal":
            abstract_num_id = abstract_num.get(qn("w:abstractNumId"))
            break

    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    number.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    number.append(level_override)
    numbering.append(number)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(number)
    p_pr.append(num_pr)


def add_steps(doc, items: list[tuple[str, str]]) -> None:
    num_id = create_numbering_instance(doc)
    for index, (title, detail) in enumerate(items):
        p = doc.add_paragraph(style="List Number")
        apply_numbering(p, num_id)
        r = p.add_run(f"{title}. ")
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(detail)
        set_run_font(r)
        set_keep_together(p)
        if index == 0 or index == len(items) - 2:
            p.paragraph_format.keep_with_next = True


def add_code(doc, code: str) -> None:
    p = doc.add_paragraph(style="Guide Code")
    set_paragraph_shading(p, LIGHT_GRAY)
    for index, line in enumerate(code.splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=9, color=DARK)
    set_keep_together(p)
    p.paragraph_format.keep_with_next = True


def add_callout(doc, label: str, text: str, kind="info") -> None:
    palettes = {
        "info": (PALE_TEAL, TEAL),
        "warning": (PALE_GOLD, GOLD),
        "danger": (PALE_RED, RED),
        "success": ("E9F5ED", GREEN),
    }
    fill, accent = palettes[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size=9)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=NAVY, size=6)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_border(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=DARK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph(style="Guide Caption")
    p.add_run(text)
    set_keep_together(p)


def add_picture(doc, path: Path, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    add_caption(doc, caption)


def font(size, bold=False):
    preferred = [
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    if bold:
        preferred = [Path("C:/Windows/Fonts/calibrib.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")]
    for candidate in preferred:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, title, subtitle="", title_color=f"#{WHITE}", subtitle_color=f"#{WHITE}"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    title_font = font(25, bold=True)
    sub_font = font(18)
    draw.text(((x1 + x2) / 2, y1 + 28), title, font=title_font, fill=title_color, anchor="mm")
    if subtitle:
        draw.multiline_text(((x1 + x2) / 2, y1 + 66), subtitle, font=sub_font, fill=subtitle_color, anchor="ma", align="center", spacing=3)


def arrow(draw, start, end, color=TEAL, width=5):
    color_rgb = f"#{color}"
    draw.line([start, end], fill=color_rgb, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(points, fill=color_rgb)


def build_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1400, 760), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (40, 90, 300, 230), f"#{NAVY}", f"#{NAVY}", "Usuarios web", "Abogado · Juez · Notario\nGestión · Auditoría · TI")
    rounded_box(draw, (40, 510, 300, 650), f"#{TEAL}", f"#{TEAL}", "Aplicación móvil", "Partes y testigos")
    rounded_box(draw, (410, 180, 670, 340), f"#{DARK_TEAL}", f"#{DARK_TEAL}", "Web Flask", "Sesión del servidor\nCSRF · navegación por perfil")
    rounded_box(draw, (780, 260, 1040, 420), f"#{NAVY}", f"#{NAVY}", "API Node.js", "JWT · RBAC · BOLA/IDOR\nreglas y auditoría")
    rounded_box(draw, (1130, 100, 1360, 250), f"#{TEAL}", f"#{TEAL}", "MySQL", "Metadatos e historial")
    rounded_box(draw, (1130, 470, 1360, 620), f"#{TEAL}", f"#{TEAL}", "MinIO", "Originales cifrados")
    arrow(draw, (300, 160), (410, 240))
    arrow(draw, (670, 260), (780, 320))
    arrow(draw, (300, 580), (780, 380))
    arrow(draw, (1040, 320), (1130, 190))
    arrow(draw, (1040, 380), (1130, 540))
    note_font = font(19, bold=True)
    draw.text((700, 710), "Toda operación jurídica pasa por la API; ningún cliente accede directamente a MySQL o MinIO.", font=note_font, fill=f"#{NAVY}", anchor="ms")
    image.save(path)


def build_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1500, 780), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    boxes = [
        ("1", "Carga", "Archivo + tipo\ndocumental"),
        ("2", "Versión", "Hash, origen y\npolítica congelada"),
        ("3", "Revisión", "Observaciones\ny respuesta"),
        ("4", "Autorización", "Solo si la regla\nlo exige"),
        ("5", "Firma", "Solo firmante\nautorizado"),
        ("6", "Certificación", "Solo si la regla\nlo exige"),
    ]
    x_positions = [35, 280, 525, 770, 1015, 1260]
    for idx, ((number, title, subtitle), x) in enumerate(zip(boxes, x_positions)):
        fill = NAVY if idx in (0, 1, 2) else TEAL
        rounded_box(draw, (x, 245, x + 205, 430), f"#{fill}", f"#{fill}", f"{number}. {title}", subtitle)
        if idx < len(boxes) - 1:
            arrow(draw, (x + 205, 337), (x_positions[idx + 1], 337), color=TEAL, width=5)
    callout_font = font(20, bold=True)
    body_font = font(18)
    draw.rounded_rectangle((250, 520, 1250, 700), radius=18, fill=f"#{LIGHT_BLUE}", outline=f"#{NAVY}", width=3)
    draw.text((750, 560), "Reglas del flujo", font=callout_font, fill=f"#{NAVY}", anchor="mm")
    draw.multiline_text((750, 610), "Una observación abierta bloquea la aprobación. Los pasos opcionales no se convierten en pendientes.\nSi una decisión debe corregirse, se emite otra que sustituye a la anterior; nunca se sobrescribe.", font=body_font, fill=f"#{DARK}", anchor="ma", align="center", spacing=7)
    image.save(path)


def add_section_intro(doc, purpose: str, audience: str | None = None) -> None:
    add_callout(doc, "Objetivo", purpose, "info")
    if audience:
        add_para(doc, f"Dirigido a: {audience}", bold_lead="Dirigido a:")


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "architecture.png"
    workflow_path = ASSET_DIR / "document_workflow.png"
    build_architecture_diagram(architecture_path)
    build_workflow_diagram(workflow_path)

    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = "Guía completa de uso de Expediente Íntegro"
    core.subject = "Manual operativo del sistema de gestión documental jurídica"
    core.author = "Equipo del Proyecto Integrador"
    core.keywords = "Expediente Íntegro, guía de usuario, gestión documental, API, Flask, MinIO"
    core.comments = "Documento de desarrollo y pruebas; no contiene secretos de producción."

    add_cover(doc)

    add_heading(doc, "Cómo utilizar esta guía", 1)
    add_section_intro(doc, "Permitir que cada perfil encuentre rápidamente sus tareas y comprenda las reglas que protegen el expediente.")
    add_para(doc, "Expediente Íntegro es una plataforma general de gestión documental jurídica. Puede utilizarse con asuntos familiares, civiles, penales, mercantiles, laborales, administrativos, constitucionales y con nuevos tipos configurables. No está limitada a nulidad matrimonial.")
    add_bullets(doc, [
        "Si eres usuario final, comienza por los capítulos 2, 4 y el capítulo correspondiente a tu espacio de trabajo.",
        "Si administras la demostración, revisa primero los capítulos 3, 10, 14 y 15.",
        "Si vas a probar un flujo completo, utiliza los escenarios del capítulo 13 y la lista de aceptación del capítulo 16.",
        "Los nombres de botones y secciones aparecen en negritas o entre comillas tal como se muestran en la interfaz.",
    ])
    add_callout(doc, "Regla de oro", "Nada se elimina físicamente. Los cambios se registran mediante nuevas versiones, estados lógicos, revocaciones, sustituciones y eventos de auditoría.", "warning")
    add_callout(doc, "Alcance actual", "La plataforma está completa para desarrollo y pruebas funcionales. La firma HMAC es una constancia de integridad interna; no sustituye una firma electrónica cualificada. El OCR y la detección automática de firmas están preparados como integración futura.", "info")

    doc.add_page_break()
    add_heading(doc, "Mapa de la guía", 2)
    add_table(doc, ["Capítulo", "Contenido"], [
        ["1", "Qué es el sistema y cómo se conectan sus componentes"],
        ["2", "Acceso, inicio de sesión y navegación común"],
        ["3", "Arranque local con Docker y verificación de servicios"],
        ["4", "Perfiles, permisos y espacios de trabajo"],
        ["5", "Expedientes, asignaciones y visibilidad"],
        ["6", "Carga, consulta y versionamiento de documentos"],
        ["7", "Revisión, observaciones, autorización, firma y certificación"],
        ["8", "Decisiones procesales inmutables"],
        ["9", "Gestión procesal: personas, etapas, plazos y estados"],
        ["10", "Auditoría y administración técnica"],
        ["11", "Uso de la aplicación móvil"],
        ["12", "Tipos documentales y reglas de análisis/firma"],
        ["13", "Escenarios completos de uso"],
        ["14", "Solución de problemas"],
        ["15", "Operación segura y mantenimiento"],
        ["16-18", "Pruebas de aceptación, credenciales de desarrollo y glosario"],
    ], [1700, 7660], font_size=9.4)

    doc.add_page_break()
    add_heading(doc, "1. Descripción general del sistema", 1)
    add_section_intro(doc, "Comprender el propósito, la arquitectura y los principios que se aplican en todas las pantallas.")
    add_heading(doc, "1.1 Propósito", 2)
    add_para(doc, "El sistema organiza expedientes jurídicos, sus participantes y sus documentos. Cada acción queda vinculada con el usuario, el perfil, el expediente y, cuando corresponde, la versión exacta del archivo. La interfaz se divide en espacios de trabajo para evitar que un usuario vea controles que no pertenecen a su función.")
    add_heading(doc, "1.2 Componentes", 2)
    add_picture(doc, architecture_path, "Figura 1. Arquitectura lógica y límites de acceso del sistema.")
    add_bullets(doc, [
        "Web Flask: presenta las pantallas para profesionales y conserva el JWT dentro de la sesión del servidor.",
        "API Node.js: autentica, aplica permisos, valida asignaciones y registra auditoría.",
        "MySQL: almacena expedientes, perfiles, políticas, estados, hashes e historial permanente.",
        "MinIO: conserva los originales cifrados y versionados en almacenamiento privado.",
        "Aplicación móvil: canal exclusivo para partes y testigos, con acceso estrictamente personal.",
    ])
    add_heading(doc, "1.3 Principios que siempre se cumplen", 2)
    add_bullets(doc, [
        "Acceso por capacidad: la interfaz muestra controles por permiso, pero la API vuelve a autorizar cada operación.",
        "Acceso por recurso: conocer el identificador de un expediente o documento no concede acceso (protección BOLA/IDOR).",
        "Versionamiento: corregir un archivo significa crear otra versión; nunca reemplazar la anterior.",
        "Inmutabilidad: las decisiones, firmas, versiones y eventos importantes no se editan ni se borran.",
        "Trazabilidad: apertura, descarga autorizada, cambio de estado y acciones documentales generan evidencia de auditoría.",
        "Política por tipo documental: un archivo solo pide autorización, firma, certificación o análisis cuando su regla lo requiere.",
    ])

    add_heading(doc, "2. Acceso y navegación", 1)
    add_section_intro(doc, "Iniciar sesión de forma segura y reconocer los elementos comunes de la web.", "Todos los perfiles web")
    add_heading(doc, "2.1 Dirección de acceso", 2)
    add_para(doc, "En el entorno local abre http://127.0.0.1:5000/login. La documentación interactiva de la API está disponible en http://127.0.0.1:3000/docs.")
    add_callout(doc, "Importante", "Las partes y los testigos no utilizan la web; deben acceder mediante la aplicación móvil.", "warning")
    add_heading(doc, "2.2 Iniciar sesión", 2)
    add_steps(doc, [
        ("Abre el portal", "Escribe la dirección de la web en un navegador actualizado."),
        ("Captura el correo", "Utiliza la cuenta asignada por Administración TI."),
        ("Captura la contraseña", "En desarrollo todas las cuentas de demostración usan 2318; en producción deben ser únicas y robustas."),
        ("Selecciona Iniciar sesión", "El sistema valida cuenta, estado, rol y canal autorizado."),
        ("Confirma el espacio", "El encabezado muestra tu nombre, perfil y espacio de trabajo. Si tienes varios perfiles, podrás entrar a los espacios permitidos."),
    ])
    add_heading(doc, "2.3 Elementos comunes", 2)
    add_table(doc, ["Elemento", "Uso"], [
        ["Encabezado", "Muestra identidad, perfil y botón Salir."],
        ["Navegación", "Cambia entre Inicio, expedientes y módulos del espacio actual."],
        ["Tarjetas de resumen", "Presentan cantidades reales obtenidas desde la API."],
        ["Estado", "Indica si el recurso está activo, pausado, cerrado, archivado u oculto."],
        ["Folio", "Identificador legible del expediente; no sustituye los controles de autorización."],
        ["Avisos", "Informan éxito, validaciones o bloqueos de la API."],
    ], [2200, 7160])
    add_heading(doc, "2.4 Cerrar sesión", 2)
    add_para(doc, "Selecciona Salir en el encabezado. El servidor invalida la sesión web y elimina el token del contexto de navegación. Cierra sesión al terminar, especialmente en equipos compartidos.")

    add_heading(doc, "3. Arranque local con Docker", 1)
    add_section_intro(doc, "Levantar y comprobar el entorno completo sin exponer MySQL al equipo anfitrión.", "Equipo de desarrollo o Administración TI")
    add_heading(doc, "3.1 Archivos de entorno", 2)
    add_para(doc, "En sistema_nulidad_api deben existir un archivo .env y un archivo .env.storage. No deben subirse a Git ni almacenarse en una ubicación compartida. Si el proyecto está en OneDrive, es preferible apuntar STORAGE_ENV_FILE a un archivo local fuera de la nube.")
    add_callout(doc, "Variable obligatoria", "El archivo .env debe incluir PLATFORM_SIGNATURE_SECRET con al menos 32 caracteres. Esta clave es independiente de MYSQL_PASSWORD, JWT_SECRET y API_APPLICATION_KEY.", "warning")
    add_code(doc, "PLATFORM_SIGNATURE_SECRET=firma_integridad_plataforma_pi_2026_local_segura")
    add_para(doc, "Para esta demostración se utiliza 2318 en los campos de contraseña solicitados por el equipo. Esta práctica solo es aceptable con NODE_ENV=development.")
    add_heading(doc, "3.2 Levantar los servicios", 2)
    add_code(doc, "cd C:\\Users\\britz\\OneDrive\\Documentos\\GitHub\\Proyecto\\sistema_nulidad_api\ndocker compose up --build -d\ndocker compose ps")
    add_para(doc, "El resultado esperado muestra api, db y minio como healthy; web debe aparecer como Up. El proceso minio-setup termina con estado Exited (0) porque es una tarea de inicialización, no un servicio permanente.")
    add_heading(doc, "3.3 Comprobar salud", 2)
    add_code(doc, "Invoke-RestMethod http://127.0.0.1:3000/health")
    add_para(doc, "La respuesta correcta indica status: ok, database: reachable y storage: reachable. Después abre la web y Swagger.")
    add_heading(doc, "3.4 Detener sin perder datos", 2)
    add_code(doc, "docker compose down")
    add_callout(doc, "Acción destructiva", "No ejecutes docker compose down -v salvo que quieras reiniciar por completo un entorno de desarrollo. La opción -v elimina los volúmenes de MySQL y MinIO.", "danger")

    add_heading(doc, "4. Perfiles y espacios de trabajo", 1)
    add_section_intro(doc, "Identificar qué puede hacer cada perfil y en qué interfaz debe trabajar.")
    add_table(doc, ["Espacio", "Perfiles", "Responsabilidad principal"], [
        ["Carga y seguimiento", "Abogado, fiscal, defensor, perito", "Aportar archivos, crear versiones, responder observaciones y revisar estados."],
        ["Revisión y decisión", "Juez, notario", "Analizar versiones y ejecutar acciones autorizadas por la regla documental."],
        ["Gestión procesal", "Secretario, coordinador", "Crear, clasificar y organizar expedientes, responsables, etapas y plazos."],
        ["Auditoría", "Auditor", "Consultar trazabilidad en modo de solo lectura y exportar reportes."],
        ["Administración técnica", "Administrador TI", "Gestionar identidades, roles, catálogos, configuración y salud."],
        ["Consulta personal móvil", "Parte, testigo", "Consultar y aportar únicamente información personal autorizada."],
    ], [2350, 2600, 4410], font_size=8.8)
    add_heading(doc, "4.1 Separación de funciones", 2)
    add_bullets(doc, [
        "El administrador TI no recibe acceso jurídico implícito a expedientes.",
        "El auditor consulta evidencias, pero no ejecuta acciones jurídicas ni administrativas.",
        "Juez y notario comparten espacio, pero los botones dependen de permisos distintos.",
        "Los perfiles de carga comparten interfaz, aunque cada expediente sigue limitado por asignación.",
        "Partes y testigos no pueden iniciar sesión en la web.",
    ])
    add_heading(doc, "4.2 Descarga y visualización", 2)
    add_table(doc, ["Perfil", "Vista en línea", "Descarga web"], [
        ["Juez", "Sí, en expedientes asignados", "Sí"],
        ["Notario", "Sí, en expedientes asignados", "Sí"],
        ["Abogado/fiscal/defensor/perito", "Sí, en expedientes asignados", "No"],
        ["Secretario/coordinador", "Según permisos y asignación", "No por defecto"],
        ["Auditor", "Metadatos y trazabilidad", "Solo reporte CSV autorizado"],
        ["Parte/testigo", "Canal móvil limitado", "No"],
    ], [3200, 3560, 2600])

    add_heading(doc, "5. Expedientes y visibilidad por asignación", 1)
    add_section_intro(doc, "Comprender por qué cada usuario ve un conjunto diferente de asuntos.")
    add_heading(doc, "5.1 Modelo centrado en casos", 2)
    add_para(doc, "La pantalla inicial muestra expedientes, no una lista global de documentos. Al abrir un expediente se presentan sus datos generales, responsables, participantes, documentos, decisiones y, para Gestión procesal, controles de etapa y plazo.")
    add_heading(doc, "5.2 Criterios de visibilidad", 2)
    add_bullets(doc, [
        "Un profesional ve los expedientes donde tiene una asignación activa y un alcance compatible.",
        "Los expedientes activos aparecen con prioridad; los pausados o cerrados se conservan según permiso histórico.",
        "Un expediente oculto deja de aparecer para perfiles sin capacidad de visibilidad histórica.",
        "Terminar una asignación es una baja lógica: el registro permanece y se guarda el motivo.",
        "Cambiar manualmente la URL o el identificador no permite abrir un caso ajeno; la API responde con acceso denegado.",
    ])
    add_callout(doc, "Si un expediente no aparece", "Comprueba primero la asignación, el estado de la cuenta, el alcance de la asignación y la visibilidad del expediente. No intentes resolverlo modificando la URL.", "info")

    add_heading(doc, "6. Carga, consulta y versionamiento documental", 1)
    add_section_intro(doc, "Registrar documentos existentes desde la web y corregirlos sin perder versiones.", "Abogado, fiscal, defensor y perito; otros perfiles cuando su permiso lo habilite")
    add_heading(doc, "6.1 Subir un documento nuevo", 2)
    add_steps(doc, [
        ("Abre Carga y seguimiento", "Selecciona uno de tus expedientes asignados."),
        ("Localiza Agregar documento", "El formulario solo aparece cuando tu sesión posee document.create."),
        ("Selecciona el tipo documental", "Este paso determina autorización, certificación, firma y análisis; no elijas un tipo aproximado."),
        ("Escribe título y descripción", "Utiliza un nombre comprensible y contexto suficiente para revisión."),
        ("Selecciona el archivo", "La web acepta PDF, JPG, JPEG o PNG válidos. El tipo se detecta por contenido, no solo por extensión."),
        ("Guarda el documento", "La API calcula SHA-256, comprime si conviene, cifra con AES-256-GCM y almacena el objeto en MinIO."),
        ("Confirma el resultado", "Abre el documento y comprueba versión 1, origen web_file, tamaño, hash y estado available."),
    ])
    add_callout(doc, "Límite de archivo", "El valor de demostración es 25 MiB (26214400 bytes). Un administrador puede cambiar MAX_DOCUMENT_BYTES para otro entorno.", "info")
    add_heading(doc, "6.2 Crear una nueva versión", 2)
    add_steps(doc, [
        ("Abre el documento", "Selecciona Consultar versiones desde la tarjeta del expediente."),
        ("Revisa observaciones", "Comprende la corrección solicitada antes de cargar otro archivo."),
        ("Selecciona Agregar una nueva versión", "Carga el archivo corregido; el documento conserva su identidad y crea un número de versión superior."),
        ("Verifica la política", "Cada versión guarda una copia inmutable de las reglas vigentes al momento de la carga."),
        ("Responde la observación", "Describe qué cambió y vincula la versión corregida cuando corresponda."),
    ])
    add_callout(doc, "Nunca sobrescribas", "No existe una acción para reemplazar versión 1. Si el archivo cambia, debe aparecer como versión 2, 3, etc.", "warning")
    add_heading(doc, "6.3 Visualizar archivos", 2)
    add_bullets(doc, [
        "Selecciona Ver en esta página en la versión deseada.",
        "PDF.js muestra PDFs dentro de la misma página y permite cambiar página y zoom.",
        "JPG y PNG se muestran como imagen protegida.",
        "Los documentos DOCX no se admiten. Convierte el archivo a PDF antes de registrarlo en el sistema.",
        "El botón Descargar original solo aparece a Juez o Notario con document.download.web.",
        "La apertura y cualquier descarga autorizada generan eventos separados de auditoría.",
    ])

    add_heading(doc, "7. Revisión y control documental", 1)
    add_section_intro(doc, "Aplicar observaciones, revisiones, autorizaciones, firmas y certificaciones en el orden permitido.", "Juez y notario")
    add_picture(doc, workflow_path, "Figura 2. Flujo documental condicionado por el tipo y la versión.")
    add_heading(doc, "7.1 Revisar una versión", 2)
    add_steps(doc, [
        ("Abre Revisión y decisión", "La bandeja prioriza expedientes activos asignados."),
        ("Abre el expediente y el documento", "Confirma que trabajas sobre la versión correcta."),
        ("Consulta el archivo", "Usa Ver en esta página y revisa los estados de análisis, autorización, certificación y firma."),
        ("Emite observación si es necesario", "Selecciona tipo, redacta una indicación precisa y guarda."),
        ("Registra la revisión", "Elige Aprobar, Devolver para corrección o Rechazar y escribe una justificación."),
    ])
    add_callout(doc, "Bloqueo automático", "Una observación abierta impide aprobar, autorizar, firmar o certificar. Devolver para corrección sí está permitido para mantener el flujo operativo.", "warning")
    add_heading(doc, "7.2 Acción del juez", 2)
    add_bullets(doc, [
        "Puede revisar documentos asignados y emitir observaciones.",
        "Puede registrar una constancia de integridad cuando la política de la versión exige firma y el rol judge está autorizado.",
        "Antes de firmar debe confirmar expresamente que revisó la versión.",
        "Puede emitir decisiones procesales sobre expedientes asignados.",
    ])
    add_heading(doc, "7.3 Acción del notario", 2)
    add_bullets(doc, [
        "Puede revisar y emitir observaciones en los expedientes asignados.",
        "Puede autorizar cuando requires_authorization está activo y la revisión previa fue aprobada.",
        "Puede certificar cuando requires_certification está activo y se cumplieron los pasos previos exigidos.",
        "No debe certificar tipos cuya política indica No requerida; la API responde con conflicto de flujo.",
    ])
    add_heading(doc, "7.4 Interpretar estados", 2)
    add_table(doc, ["Estado", "Significado operativo"], [
        ["Pendiente", "Todavía falta ejecutar la acción obligatoria."],
        ["Aprobado / Autorizado / Firmado", "La acción quedó registrada para esa versión exacta."],
        ["Devuelto", "Se requiere corrección; normalmente se crea una versión nueva."],
        ["Rechazado", "La versión no fue aceptada en ese paso."],
        ["No requerido", "La política del tipo no exige esa acción; no bloquea el flujo."],
        ["Opcional", "Puede ejecutarse si hay justificación, pero su ausencia no bloquea."],
    ], [2600, 6760])
    add_callout(doc, "Firma interna", "La constancia HMAC-SHA256 vincula versión, usuario, rol y momento. Sirve para integridad y trazabilidad interna; no equivale por sí sola a e.firma o a un certificado legal cualificado.", "info")

    add_heading(doc, "8. Decisiones procesales inmutables", 1)
    add_section_intro(doc, "Emitir determinaciones y corregirlas sin modificar el registro original.", "Juez con permiso case.decision.issue")
    add_heading(doc, "8.1 Emitir una decisión", 2)
    add_steps(doc, [
        ("Abre el expediente", "Confirma folio, materia, tipo de asunto y estado."),
        ("Selecciona el tipo", "Acuerdo procesal, resolución interlocutoria, sentencia final u otra determinación."),
        ("Escribe el resultado", "Usa un código o expresión breve y consistente, por ejemplo procedente o admite_prueba."),
        ("Redacta título y fundamentación", "La fundamentación debe ser suficiente porque no podrá editarse."),
        ("Vincula una versión", "La sentencia final exige una versión documental que haya completado su flujo."),
        ("Emite la decisión", "La API confirma asignación, estado y requisitos antes de insertar el registro."),
    ])
    add_heading(doc, "8.2 Corregir una decisión", 2)
    add_para(doc, "No se modifica la decisión anterior. Crea otra decisión y selecciona el registro en Sustituye a. La nueva queda vigente y la anterior se muestra como Sustituida, conservando fechas, autor y contenido original.")
    add_callout(doc, "Protección comprobada", "Los triggers de MySQL rechazan UPDATE o DELETE sobre decisiones. Si un texto histórico contiene un error, la corrección correcta es una decisión sustitutiva.", "success")

    add_heading(doc, "9. Gestión procesal", 1)
    add_section_intro(doc, "Crear y organizar expedientes sin borrar asignaciones, participantes, plazos o etapas.", "Secretario y coordinador")
    add_heading(doc, "9.1 Crear un expediente", 2)
    add_steps(doc, [
        ("Abre Gestión procesal > Expedientes", "Selecciona el formulario Crear expediente."),
        ("Captura folio, título y descripción", "El folio debe ser único y seguir la convención institucional."),
        ("Clasifica el asunto", "Elige tipo de expediente y unidad organizacional; la materia se obtiene del catálogo."),
        ("Define confidencialidad", "Interna, confidencial o restringida según el caso."),
        ("Selecciona estado inicial", "Activo para operación inmediata o Borrador para preparación."),
        ("Registra el motivo", "Todo alta y cambio de estado necesita una explicación auditable."),
    ])
    add_heading(doc, "9.2 Asignar responsables", 2)
    add_bullets(doc, [
        "Selecciona usuario, tipo de asignación y alcance: completo, documental, procesal, revisión o auditoría.",
        "Escribe el motivo de la asignación.",
        "Para retirar acceso, utiliza Finalizar asignación; no elimines la fila.",
        "Una asignación finalizada permanece en el historial con actor, fecha y motivo.",
    ])
    add_heading(doc, "9.3 Agregar participantes", 2)
    add_para(doc, "Selecciona una persona y su calidad procesal. La aplicación permite finalizar la participación de forma lógica. Los datos personales y documentos propios siguen sujetos a permisos de recurso.")
    add_heading(doc, "9.4 Cambiar etapa", 2)
    add_steps(doc, [
        ("Revisa la etapa actual", "Aparece en el resumen del expediente."),
        ("Selecciona una etapa disponible", "Las opciones provienen de la definición del tipo de asunto."),
        ("Escribe el motivo", "Explica el hito que justifica la transición."),
        ("Registra la transición", "El sistema conserva la etapa anterior y la nueva en el historial."),
    ])
    add_heading(doc, "9.5 Crear y cerrar plazos", 2)
    add_bullets(doc, [
        "Captura título, vencimiento, responsable opcional, descripción y motivo.",
        "Un plazo pendiente puede terminar como Completado, Cancelado o Vencido.",
        "La identidad del plazo es inmutable; si fue creado incorrectamente, finalízalo con motivo y crea otro.",
        "Los plazos próximos alimentan los indicadores de Gestión procesal.",
    ])
    add_heading(doc, "9.6 Estado y visibilidad", 2)
    add_table(doc, ["Acción", "Efecto"], [
        ["Pausar", "Conserva el expediente y limita operaciones hasta reanudación."],
        ["Cerrar", "Finaliza la operación ordinaria sin eliminar documentos ni historial."],
        ["Anular", "Registra anulación lógica y motivo; los perfiles históricos autorizados conservan trazabilidad."],
        ["Archivar", "Mueve el asunto a conservación de largo plazo según política."],
        ["Ocultar", "Impide visibilidad ordinaria; no borra el registro ni el almacenamiento."],
    ], [2300, 7060])

    add_heading(doc, "10. Auditoría y administración técnica", 1)
    add_section_intro(doc, "Consultar evidencia de trazabilidad y mantener la plataforma sin mezclar funciones jurídicas y técnicas.")
    add_heading(doc, "10.1 Espacio de Auditoría", 2)
    add_bullets(doc, [
        "Inicio: resumen de eventos, versiones y firmas consultables.",
        "Eventos: acciones registradas con fecha, actor, acción y recurso.",
        "Accesos: aperturas y descargas autorizadas.",
        "Versiones: historial documental sin eliminación.",
        "Firmas: constancias de integridad y datos de verificación.",
        "Reportes: exportación CSV de hasta 5000 eventos en modo de solo lectura.",
    ])
    add_callout(doc, "Solo lectura", "El auditor no puede modificar expedientes, documentos, usuarios ni catálogos. El reporte es una copia; la bitácora original no cambia.", "info")
    add_heading(doc, "10.2 Usuarios", 2)
    add_steps(doc, [
        ("Abre Administración técnica > Usuarios", "Consulta cuentas, roles y estado."),
        ("Crea la cuenta", "Captura nombre, correo, contraseña inicial, rol y motivo."),
        ("Concede o revoca roles", "Cada concesión queda como registro independiente y revocable."),
        ("Cambia el estado", "Usa Activo, Suspendido o Archivado; nunca elimines el usuario."),
    ])
    add_heading(doc, "10.3 Catálogos y tipos documentales", 2)
    add_para(doc, "Administración TI puede crear tipos documentales y su primera regla. Esa regla queda inmutable para proteger versiones ya cargadas. Una modificación futura debe publicarse como una regla nueva, no reescribir la anterior.")
    add_heading(doc, "10.4 Configuración y salud", 2)
    add_bullets(doc, [
        "Configuración muestra valores no sensibles necesarios para diagnosticar el entorno.",
        "Salud confirma conectividad con API, base de datos y almacenamiento.",
        "La interfaz nunca debe revelar JWT_SECRET, API_APPLICATION_KEY, claves de cifrado o contraseñas.",
        "Administración TI no obtiene permiso automático para leer documentos jurídicos.",
    ])

    add_heading(doc, "11. Aplicación móvil para partes y testigos", 1)
    add_section_intro(doc, "Diferenciar el canal personal móvil del trabajo profesional en la web.", "Partes y testigos")
    add_bullets(doc, [
        "La persona inicia sesión con una cuenta cuyo canal es mobile.",
        "Solo consulta expedientes donde participa y documentos que le pertenecen.",
        "Puede aportar archivos propios y, cuando la función esté habilitada, capturar con cámara.",
        "No puede descargar originales ni acceder a documentos de otras personas.",
        "No puede autorizar, certificar, firmar, administrar usuarios ni emitir decisiones.",
        "La API vuelve a validar propiedad y participación en cada solicitud.",
    ])
    add_callout(doc, "Separación de canales", "Un juez, notario o abogado debe usar la web. Una parte o testigo debe usar la aplicación móvil. El canal solicitado durante el login debe coincidir con un rol activo.", "warning")

    add_heading(doc, "12. Tipos documentales y reglas", 1)
    add_section_intro(doc, "Configurar correctamente cuándo un archivo necesita análisis, autorización, firma o certificación.", "Administración TI y responsables de política documental")
    add_table(doc, ["Campo", "Cómo decidirlo"], [
        ["Código", "Identificador estable en minúsculas, números y guion bajo; no debe reciclarse."],
        ["Propietario", "case para documentos del expediente; participant para documentos personales."],
        ["Datos sensibles", "Actívalo para identidad, domicilio, datos fiscales, pruebas reservadas o información personal."],
        ["Firma incorporada", "Indica si el archivo debe contener una firma visible/digital antes de subirlo."],
        ["Firma de plataforma", "Indica si un rol debe registrar una constancia interna sobre esa versión."],
        ["Autorización", "Actívala cuando se necesita visto bueno formal antes de continuar."],
        ["Certificación", "Actívala cuando el notario deba certificar el resultado final del flujo."],
        ["Análisis", "skip, on_demand o automatic según el futuro servicio OCR/detección de firmas."],
        ["Roles firmantes", "Lista exacta de perfiles autorizados para la constancia de plataforma."],
    ], [2650, 6710], font_size=8.9)
    add_heading(doc, "12.1 Buena práctica recomendada", 2)
    add_para(doc, "Nunca envíes todos los archivos al mismo verificador de firma. El analizador debe recibir la política congelada de la versión y responder en el contexto del tipo documental. De esta forma, una identificación o un anexo que no requiere firma no queda bloqueado por un falso pendiente.")
    add_heading(doc, "12.2 Integración futura del analizador", 2)
    add_steps(doc, [
        ("Crear trabajo", "La API registra document_analysis_jobs con versión, política y proveedor."),
        ("Analizar contenido", "El servicio externo ejecuta OCR, detección de firma u otras verificaciones autorizadas."),
        ("Guardar resultado", "document_analysis_results conserva hallazgos, confianza y evidencia."),
        ("Aplicar la regla", "La API decide si el hallazgo es informativo o bloqueante según la política de la versión."),
        ("Auditar", "Solicitud, respuesta y decisión quedan ligadas a la versión; el analizador no modifica el original."),
    ])

    add_heading(doc, "13. Escenarios completos de uso", 1)
    add_section_intro(doc, "Probar el sistema siguiendo secuencias realistas y verificables.")
    add_heading(doc, "13.1 Caso civil con documento corregido", 2)
    add_steps(doc, [
        ("Gestión procesal", "El secretario crea el expediente civil, asigna abogado y juez, y agrega participantes."),
        ("Carga", "El abogado selecciona el caso, elige el tipo documental correcto y carga versión 1."),
        ("Revisión", "El juez abre versión 1 y emite una observación de corrección."),
        ("Corrección", "El abogado carga versión 2 y responde la observación vinculando la nueva versión."),
        ("Resolución", "El revisor marca la observación como resuelta y aprueba versión 2."),
        ("Firma", "Si la regla lo exige, el juez registra la constancia de integridad."),
        ("Decisión", "El juez emite la decisión y vincula la versión que completó el flujo."),
        ("Auditoría", "El auditor comprueba cargas, visualizaciones, observación, respuesta, firma y decisión."),
    ])
    add_heading(doc, "13.2 Documento que no requiere firma", 2)
    add_steps(doc, [
        ("Carga", "Selecciona un tipo cuya firma de plataforma sea none u optional."),
        ("Consulta", "Comprueba que el estado muestre No requerida u Opcional, no Pendiente."),
        ("Prueba negativa", "Intenta certificar o firmar con un rol no autorizado; la API debe devolver un conflicto y no crear registros."),
        ("Auditoría", "Verifica que el documento continúa íntegro y que no se alteró su política."),
    ])
    add_heading(doc, "13.3 Baja lógica", 2)
    add_steps(doc, [
        ("Selecciona un recurso de prueba", "Utiliza una cuenta o expediente creado para QA."),
        ("Cambia el estado", "Suspende/archiva la cuenta u oculta/cierra el expediente e indica motivo."),
        ("Comprueba visibilidad", "El recurso deja de aparecer para perfiles ordinarios."),
        ("Comprueba conservación", "El auditor o administrador autorizado todavía encuentra el historial y la acción."),
    ])

    add_heading(doc, "14. Solución de problemas", 1)
    add_section_intro(doc, "Diagnosticar fallos comunes sin eliminar volúmenes ni debilitar los controles de seguridad.")
    add_table(doc, ["Síntoma", "Causa probable", "Qué hacer"], [
        ["MYSQL_DATABASE is missing", "Compose no leyó el .env correcto.", "Ejecuta desde sistema_nulidad_api, confirma el nombre exacto .env y usa docker compose config."],
        [".env.storage not found", "STORAGE_ENV_FILE apunta a una ruta inexistente.", "Crea el archivo desde .env.storage.example o corrige la ruta absoluta/local."],
        ["db is unhealthy", "Contraseña inicial faltante, volumen inconsistente o MySQL aún inicia.", "Revisa docker compose logs db y las variables MYSQL_*; no borres el volumen sin respaldo."],
        ["Web reinicia o no abre", "API no saludable, secreto Flask ausente o error de plantilla.", "Revisa docker compose ps y docker compose logs web api."],
        ["403 origin_not_allowed", "Swagger/origen no aparece en CORS_ORIGINS.", "Incluye http://127.0.0.1:3000 y reinicia la API."],
        ["403 al abrir un caso", "No existe asignación activa o el canal/rol no coincide.", "Pide a Gestión procesal revisar la asignación; no cambies el ID de la URL."],
        ["El caso no aparece", "Está oculto, cerrado o asignado a otra persona.", "Comprueba estado, visibilidad, alcance y fecha de fin de asignación."],
        ["No aparece Descargar", "El perfil carece de document.download.web.", "Es comportamiento esperado; solo juez/notario autorizados descargan."],
        ["DOCX rechazado", "Word no forma parte de los formatos permitidos.", "Convierte el documento a PDF antes de cargarlo."],
        ["409 al firmar/certificar", "Política no aplicable, observación abierta o paso previo pendiente.", "Lee los estados de la versión y completa el flujo en orden."],
        ["Archivo rechazado", "El contenido no es PDF/JPG/PNG válido o supera el límite.", "Verifica el archivo real, no solo su extensión, y revisa MAX_DOCUMENT_BYTES."],
        ["MinIO no disponible", "Servicio, credenciales o bucket no están listos.", "Revisa logs minio y minio-setup; no publiques el bucket como público."],
    ], [2350, 2950, 4060], font_size=8.0)
    add_heading(doc, "14.1 Comandos de diagnóstico", 2)
    add_code(doc, "docker compose ps\ndocker compose logs --tail 100 api\ndocker compose logs --tail 100 web\ndocker compose logs --tail 100 db\ndocker compose logs --tail 100 minio\nInvoke-RestMethod http://127.0.0.1:3000/health")
    add_callout(doc, "Evita la pérdida de evidencia", "No uses DELETE, DROP, TRUNCATE, docker compose down -v ni limpieza manual de MinIO como respuesta inicial a un error.", "danger")

    add_heading(doc, "15. Operación segura y mantenimiento", 1)
    add_section_intro(doc, "Mantener el entorno disponible y proteger secretos, originales e historial.", "Administración TI y desarrollo")
    add_heading(doc, "15.1 Secretos", 2)
    add_bullets(doc, [
        "Usa claves distintas para MySQL, JWT, API de sistema, firma de plataforma, sesión Flask, MinIO y cifrado documental.",
        "No copies secretos a JavaScript del navegador ni a variables EXPO_PUBLIC_*.",
        "No subas .env, .env.storage, certificados privados, respaldos ni documentos a Git.",
        "En producción reemplaza 2318 por contraseñas únicas y aleatorias; establece NODE_ENV=production.",
        "Usa HTTPS y SESSION_COOKIE_SECURE=true antes de exponer la web fuera de localhost.",
    ])
    add_heading(doc, "15.2 Respaldo y recuperación", 2)
    add_bullets(doc, [
        "Realiza respaldos coherentes de MySQL y del volumen de MinIO; uno sin el otro deja metadatos u objetos incompletos.",
        "Cifra respaldos y limita su acceso; conserva las claves fuera del mismo servidor.",
        "Prueba la restauración en un entorno aislado y documenta tiempos de recuperación.",
        "No confíes en la retención de MinIO como sustituto de una estrategia de respaldo.",
        "Antes de producción define monitoreo, alertas, rotación de secretos y política institucional de conservación.",
    ])
    add_heading(doc, "15.3 Actualizaciones", 2)
    add_steps(doc, [
        ("Respalda", "Obtén respaldo verificado de base y almacenamiento."),
        ("Revisa migraciones", "Confirma que sean aditivas y respeten triggers de inmutabilidad."),
        ("Reconstruye", "Ejecuta docker compose up --build -d."),
        ("Valida", "Revisa salud, login y un flujo representativo por espacio."),
        ("Audita", "Comprueba que los registros anteriores permanecen disponibles."),
    ])

    add_heading(doc, "16. Lista de aceptación funcional", 1)
    add_section_intro(doc, "Confirmar que el sistema está listo para una demostración o sesión de pruebas.")
    acceptance = [
        ["Acceso", "Cada cuenta web entra a su espacio y parte/testigo son rechazados por la web.", "☐"],
        ["Asignación", "Un usuario solo ve casos asignados; un ID ajeno devuelve 403/404 controlado.", "☐"],
        ["Carga", "Un PDF/JPG/PNG válido crea el documento y su versión en MinIO.", "☐"],
        ["Versiones", "La corrección crea una versión superior y conserva la anterior.", "☐"],
        ["Visor", "PDF e imagen abren en la misma página sin botón de descarga para perfil restringido.", "☐"],
        ["Política", "Un tipo sin firma muestra No requerida y no bloquea el flujo.", "☐"],
        ["Observación", "Una observación abierta bloquea aprobación/firma; respuesta y resolución quedan registradas.", "☐"],
        ["Revisión", "Juez/notario ejecutan únicamente acciones de su permiso y del tipo documental.", "☐"],
        ["Decisión", "Una corrección crea decisión sustitutiva; la anterior queda como Sustituida.", "☐"],
        ["Proceso", "Se crean expediente, asignación, participante, transición y plazo con motivo.", "☐"],
        ["Baja lógica", "Suspender/archivar/ocultar no elimina registros ni archivos.", "☐"],
        ["Auditoría", "Eventos, accesos, versiones, firmas y CSV reflejan las acciones realizadas.", "☐"],
        ["Administración", "Usuario, rol y tipo documental se gestionan sin conceder acceso legal al perfil TI.", "☐"],
        ["Salud", "API, MySQL y MinIO aparecen reachable/healthy y la web responde.", "☐"],
    ]
    add_table(doc, ["Área", "Criterio", "OK"], acceptance, [1800, 6760, 800], font_size=8.6)

    add_heading(doc, "17. Credenciales de desarrollo", 1)
    add_section_intro(doc, "Probar cada espacio con cuentas sintéticas. Estas credenciales no deben utilizarse en producción.")
    add_callout(doc, "Contraseña común de desarrollo", "2318", "warning")
    credentials = [
        ["Administración TI", "admin.web@example.test", "Administración técnica", "web/technical"],
        ["Abogado", "web.abogado.20260731@example.test", "Carga y seguimiento", "web"],
        ["Fiscal", "web.fiscal.20260731@example.test", "Carga y seguimiento", "web"],
        ["Defensor", "web.defensor.20260731@example.test", "Carga y seguimiento", "web"],
        ["Perito", "web.perito.20260731@example.test", "Carga y seguimiento", "web"],
        ["Juez", "web.juez.20260731@example.test", "Revisión y decisión", "web"],
        ["Notario", "web.notario.20260731@example.test", "Revisión y decisión", "web"],
        ["Secretario", "web.secretario.20260731@example.test", "Gestión procesal", "web"],
        ["Coordinador", "web.coordinador.20260731@example.test", "Gestión procesal", "web"],
        ["Auditor", "web.auditor.20260731@example.test", "Auditoría", "web"],
        ["Parte", "mobile.parte.20260731@example.test", "Consulta personal", "mobile"],
        ["Testigo", "mobile.testigo.20260731@example.test", "Consulta personal", "mobile"],
    ]
    add_table(doc, ["Perfil", "Correo", "Espacio", "Canal"], credentials, [1800, 3600, 2660, 1300], font_size=7.9)
    add_bullets(doc, [
        "Si una cuenta no entra, revisa que account_status sea active y que el rol tenga una concesión vigente.",
        "Las cuentas profesionales necesitan asignaciones para ver expedientes; iniciar sesión no concede acceso global.",
        "Después de la demostración, cambia o elimina lógicamente las credenciales de prueba antes de usar datos reales.",
    ])

    add_heading(doc, "18. Glosario", 1)
    add_section_intro(doc, "Unificar el significado de los términos que aparecen en la interfaz y en la API.")
    glossary = [
        ["Asignación", "Relación activa entre un profesional y un expediente, con tipo y alcance."],
        ["Autorización", "Visto bueno formal exigido por la política antes de acciones posteriores."],
        ["BOLA/IDOR", "Falla evitada al verificar que el usuario puede acceder al recurso solicitado por su identificador."],
        ["Baja lógica", "Cambio de estado que oculta o desactiva sin eliminar físicamente."],
        ["Certificación", "Acción notarial final cuando la política de la versión la requiere."],
        ["Constancia de integridad", "Registro HMAC ligado a la versión, actor, rol y momento; no es firma cualificada."],
        ["Expediente", "Contenedor jurídico de participantes, asignaciones, documentos, etapas, plazos y decisiones."],
        ["Firma incorporada", "Firma que debe existir dentro del archivo antes del análisis."],
        ["Hash SHA-256", "Huella que permite comprobar si los bytes del archivo cambiaron."],
        ["MinIO", "Almacenamiento privado compatible con S3 para originales cifrados y versionados."],
        ["Observación", "Solicitud de comentario, corrección, aclaración o revisión jurídica sobre una versión."],
        ["Origen", "Canal de incorporación: web_file, mobile_file, mobile_camera o legado."],
        ["Política congelada", "Copia inmutable de la regla documental aplicada a una versión al momento de crearla."],
        ["RBAC", "Control de acceso basado en roles y permisos activos."],
        ["Versión", "Instancia inmutable de un archivo dentro de la identidad de un documento."],
    ]
    add_table(doc, ["Término", "Definición"], glossary, [2400, 6960], font_size=8.8)

    add_heading(doc, "Referencia rápida", 1)
    add_section_intro(doc, "Concentrar direcciones y comandos básicos para una demostración local.")
    add_table(doc, ["Recurso", "Dirección"], [
        ["Inicio de sesión web", "http://127.0.0.1:5000/login"],
        ["Documentación Swagger", "http://127.0.0.1:3000/docs"],
        ["Contrato OpenAPI", "http://127.0.0.1:3000/openapi.json"],
        ["Salud de API", "http://127.0.0.1:3000/health"],
        ["Consola MinIO (local)", "http://127.0.0.1:9001"],
    ], [3000, 6360])
    add_code(doc, "docker compose up --build -d\ndocker compose ps\ndocker compose logs --tail 100 api\ndocker compose down")
    add_callout(doc, "Fin de la guía", "Para una prueba completa, ejecuta el capítulo 16 y conserva el reporte de Auditoría como evidencia de la sesión.", "success")

    # Keep the document update-friendly and ensure fields recalculate in Word.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    result = build_document()
    print(result)
